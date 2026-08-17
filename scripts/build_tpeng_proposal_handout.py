#!/usr/bin/env python3
"""Build the LabelLab proposal PDF and a visually faithful Word handout.

Markdown is the editable source of truth.  The Word handout intentionally
embeds the rendered PDF pages because the headless LibreOffice runtime on the
Mac can substitute missing CJK glyphs when rendering editable OOXML runs.
"""

from __future__ import annotations

import os
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


CODEX_NODE = Path(
    "/Users/yukina/.cache/codex-runtimes/codex-primary-runtime/dependencies/node/bin/node"
)
CODEX_NODE_MODULES = Path(
    "/Users/yukina/.cache/codex-runtimes/codex-primary-runtime/dependencies/node/node_modules"
)
BROWSER_CANDIDATES = (
    Path("/Applications/Microsoft Edge.app/Contents/MacOS/Microsoft Edge"),
    Path("/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"),
)


CSS = r"""
@page { size: Letter; margin: 0.72in 0.82in 0.72in 0.82in; }
* { box-sizing: border-box; }
html, body { margin: 0; padding: 0; }
body {
  font-family: "Source Han Sans CN", "PingFang SC", "Hiragino Sans GB", Arial, sans-serif;
  font-size: 11pt;
  line-height: 1.52;
  color: #182230;
  background: #fff;
}
header { border-bottom: 1px solid #c9d3e0; padding: 0 0 8px; margin-bottom: 18px; text-align: right; color: #6b7280; font-size: 9pt; }
footer { position: fixed; right: 0; bottom: -0.34in; color: #6b7280; font-size: 9pt; }
h1 { font-size: 24pt; line-height: 1.18; color: #0b2545; margin: 0 0 7pt; page-break-after: avoid; }
h2 { font-size: 16pt; line-height: 1.25; color: #2e74b5; margin: 18pt 0 7pt; page-break-after: avoid; }
h3 { font-size: 13pt; line-height: 1.25; color: #1f4d78; margin: 12pt 0 5pt; page-break-after: avoid; }
h4 { font-size: 11.5pt; color: #1f4d78; margin: 9pt 0 4pt; page-break-after: avoid; }
p { margin: 0 0 7pt; widows: 3; orphans: 3; }
strong { color: #0b2545; }
blockquote { margin: 10pt 0 12pt; padding: 9pt 12pt; border-left: 4px solid #2e74b5; background: #f4f6f9; color: #243447; }
ul, ol { margin: 4pt 0 8pt 20pt; padding-left: 13pt; }
li { margin: 0 0 4pt; }
table { width: 100%; table-layout: fixed; border-collapse: collapse; margin: 9pt 0 12pt; font-size: 8.8pt; page-break-inside: auto; }
thead { display: table-header-group; }
tr { page-break-inside: avoid; }
th, td { border: 1px solid #c9d3e0; padding: 5pt 6pt; vertical-align: top; overflow-wrap: anywhere; word-break: break-word; }
th { background: #f2f4f7; color: #0b2545; font-weight: 700; }
code { font-family: "SFMono-Regular", Menlo, monospace; font-size: 0.9em; color: #1f4d78; }
pre { margin: 8pt 0 10pt; padding: 8pt 10pt; background: #f6f8fb; border: 1px solid #d7e0ea; overflow-wrap: anywhere; white-space: pre-wrap; font-size: 8.5pt; line-height: 1.35; }
hr { border: 0; border-top: 1px solid #d7e0ea; margin: 14pt 0; }
.meta { display: grid; grid-template-columns: 1.05in 1fr; gap: 4pt 12pt; margin: 10pt 0 13pt; }
.meta b { color: #0b2545; }
.lead { margin: 7pt 0 13pt; padding: 10pt 12pt; background: #f4f6f9; border-left: 4px solid #2e74b5; }
.workflow { margin: 10pt 0 12pt; padding: 9pt 11pt; background: #eef4fb; border: 1px solid #a9bfd8; color: #0b2545; font-weight: 600; }
"""


def preprocess(markdown: str) -> str:
    """Remove Mermaid syntax while preserving its business meaning in print."""
    pattern = re.compile(r"```mermaid\s*.*?```", re.S)
    replacement = (
        "> **业务闭环：** 字段需求合同 → 素材接入 → 自动/人工标注 → 人工纠偏 → "
        "AI 候选机制与回归 → 人工启用机制 → 正式标签发布 → 下游投影消费 → Badcase 回流"
    )
    return pattern.sub(replacement, markdown)


def find_node() -> Path:
    configured = os.environ.get("TPENG_NODE")
    if configured:
        candidate = Path(configured).expanduser()
        if candidate.exists():
            return candidate
        raise RuntimeError(f"TPENG_NODE does not exist: {candidate}")
    system_node = shutil.which("node")
    candidates = [Path(system_node)] if system_node else []
    candidates.extend((CODEX_NODE, Path("/opt/homebrew/bin/node")))
    for candidate in candidates:
        if candidate.exists():
            return candidate
    raise RuntimeError("a Node.js runtime is required for PDF rendering")


def find_browser() -> Path:
    configured = os.environ.get("TPENG_BROWSER")
    if configured:
        candidate = Path(configured).expanduser()
        if candidate.exists():
            return candidate
        raise RuntimeError(f"TPENG_BROWSER does not exist: {candidate}")
    for candidate in BROWSER_CANDIDATES:
        if candidate.exists():
            return candidate
    raise RuntimeError("Microsoft Edge or Google Chrome is required for PDF rendering")


def node_path_for_playwright() -> str | None:
    configured = os.environ.get("NODE_PATH")
    if configured:
        return configured
    repo_modules = Path(__file__).resolve().parents[1] / "frontend" / "node_modules"
    for candidate in (repo_modules, CODEX_NODE_MODULES):
        if (candidate / "playwright").exists():
            return str(candidate)
    return None


def render_pdf(source: Path, pdf_path: Path, html_path: Path) -> None:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        raise RuntimeError("pandoc is required for proposal PDF rendering")
    rendered_md = preprocess(source.read_text(encoding="utf-8"))
    with tempfile.TemporaryDirectory(prefix="tpeng-proposal-html-") as tmp:
        md_path = Path(tmp) / "proposal.md"
        md_path.write_text(rendered_md, encoding="utf-8")
        raw_html = subprocess.check_output(
            [
                pandoc,
                str(md_path),
                "--from=gfm",
                "--to=html5",
                "--standalone",
                "--metadata",
                "pagetitle=TPENG 新标签体系方案",
            ],
            text=True,
        )
    html_path.write_text(
        raw_html.replace("</head>", f"<style>{CSS}</style></head>", 1), encoding="utf-8"
    )
    node = find_node()
    print_helper = Path(__file__).with_name("print_tpeng_proposal_pdf.mjs")
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    env = os.environ.copy()
    node_path = node_path_for_playwright()
    if node_path:
        env["NODE_PATH"] = node_path
    env["TPENG_BROWSER"] = str(find_browser())
    subprocess.run(
        [str(node), str(print_helper), str(html_path.resolve()), str(pdf_path.resolve())],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        env=env,
    )


def set_page_geometry(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)


def disable_paragraph_spacing(paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = 0
    fmt.space_after = 0
    fmt.line_spacing = 1


def add_full_page_image(doc: Document, image_path: Path, first: bool) -> None:
    if not first:
        page_break = doc.add_paragraph()
        disable_paragraph_spacing(page_break)
        page_break.paragraph_format.page_break_before = True
    paragraph = doc.add_paragraph()
    disable_paragraph_spacing(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    # Leave a small amount of vertical room for the paragraph mark/page-break
    # Word inserts around inline pictures; using the full 8.5in width causes
    # LibreOffice to spill the break onto a blank page.
    run.add_picture(str(image_path), width=Inches(8.25))


def build_image_handout(pdf_path: Path, output_docx: Path, render_dir: Path) -> int:
    pdftoppm = shutil.which("pdftoppm")
    if not pdftoppm:
        raise RuntimeError("pdftoppm is required to build the visual Word handout")
    render_dir.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [pdftoppm, "-png", "-r", "144", str(pdf_path), str(render_dir / "page")],
        check=True,
    )
    pages = sorted(render_dir.glob("page-*.png"), key=lambda p: int(p.stem.split("-")[-1]))
    if not pages:
        raise RuntimeError("PDF rendering produced no pages")
    doc = Document()
    set_page_geometry(doc.sections[0])
    for index, page in enumerate(pages):
        add_full_page_image(doc, page, first=index == 0)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)
    return len(pages)


def main(argv: list[str]) -> int:
    if len(argv) != 4:
        raise SystemExit("usage: build_tpeng_proposal_handout.py SOURCE.md OUTPUT.pdf OUTPUT.docx")
    source, output_pdf, output_docx = map(Path, argv[1:])
    with tempfile.TemporaryDirectory(prefix="tpeng-proposal-build-") as tmp:
        tmp_path = Path(tmp)
        html_path = tmp_path / "proposal.html"
        render_dir = tmp_path / "pages"
        render_pdf(source, output_pdf, html_path)
        page_count = build_image_handout(output_pdf, output_docx, render_dir)
    print(f"rendered {page_count} pages")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
