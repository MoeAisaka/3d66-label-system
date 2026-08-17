#!/usr/bin/env python3
"""Build a polished Word handout from the authoritative LabelLab proposal Markdown."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BODY_FONT = "Hiragino Sans GB"
LATIN_FONT = "Arial"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
BORDER = "C9D3E0"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, color=BORDER, size="4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, size=11, bold=False, color="000000", italic=False, mono=False) -> None:
    run.font.name = "Menlo" if mono else LATIN_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Menlo" if mono else LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), "Menlo" if mono else LATIN_FONT)
    rfonts.set(qn("w:cs"), BODY_FONT)
    rfonts.set(qn("w:eastAsia"), BODY_FONT)


def set_para(paragraph, before=0, after=6, line=1.10, alignment=None, left=0, first=0) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if alignment is not None:
        paragraph.alignment = alignment
    if left:
        fmt.left_indent = Inches(left)
    if first:
        fmt.first_line_indent = Inches(first)


def add_inline(paragraph, text: str, size=11, color="000000") -> None:
    # Keep the parser intentionally small and deterministic for this proposal.
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=max(9, size - 1), color=DARK_BLUE, mono=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, italic=True, color=color)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, color=color)


def add_callout(doc: Document, text: str, emphasis=False) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.25)
    set_cell_shading(cell, CALLOUT)
    set_cell_borders(cell, color=BLUE, size="10")
    set_cell_margins(cell, top=120, start=180, bottom=120, end=180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    set_para(p, before=0, after=0, line=1.20)
    add_inline(p, text, size=11, color=INK)
    if emphasis and p.runs:
        p.runs[0].bold = True
    spacer = doc.add_paragraph()
    set_para(spacer, before=0, after=2, line=1.0)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [6.25 / cols] * cols
    for ridx, row in enumerate(rows):
        for cidx in range(cols):
            cell = table.cell(ridx, cidx)
            cell.width = Inches(widths[cidx])
            set_cell_margins(cell)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            set_para(p, before=0, after=0, line=1.05)
            value = row[cidx] if cidx < len(row) else ""
            add_inline(p, value, size=9.5, color=INK if ridx == 0 else "000000")
            if ridx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
                for run in p.runs:
                    run.bold = True
    spacer = doc.add_paragraph()
    set_para(spacer, before=0, after=4, line=1.0)


def parse_table(lines: list[str]) -> list[list[str]]:
    out = []
    for line in lines:
        if not line.strip().startswith("|"):
            continue
        parts = [p.strip() for p in line.strip().strip("|").split("|")]
        if parts and all(re.fullmatch(r":?-{3,}:?", p) for p in parts):
            continue
        out.append(parts)
    return out


def page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color="6B7280")


def build(source: Path, output: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    doc.core_properties.title = "TPENG 新标签体系方案"
    doc.core_properties.subject = "LabelLab 统一底座与 45 天 MVP Roadmap"
    doc.core_properties.author = "TPENG"

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    header = section.header.paragraphs[0]
    set_para(header, before=0, after=0, line=1.0, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_inline(header, "TPENG 标签实验台（LabelLab）｜最终方案", size=9, color="6B7280")
    footer = section.footer.paragraphs[0]
    set_para(footer, before=0, after=0, line=1.0, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_inline(footer, "内部方案｜第 ", size=9, color="6B7280")
    page_field(footer)
    add_inline(footer, " 页", size=9, color="6B7280")

    lines = source.read_text(encoding="utf-8").splitlines()
    i = 0
    title_done = False
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("```"):
            fence = stripped
            code: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            if code and code[0].startswith("flowchart"):
                add_callout(doc, "业务闭环：字段需求 → 素材接入 → 自动标注 → 人工纠偏 → AI 候选回归 → 人工启用 → 正式发布 → 投影消费 → Badcase 回流", emphasis=True)
            else:
                p = doc.add_paragraph()
                set_para(p, before=3, after=8, line=1.0)
                p.paragraph_format.left_indent = Inches(0.18)
                p.paragraph_format.right_indent = Inches(0.18)
                for idx, line in enumerate(code):
                    run = p.add_run(line)
                    set_run_font(run, size=8.5, color="334155", mono=True)
                    if idx != len(code) - 1:
                        run.add_break()
            continue
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, parse_table(table_lines))
            continue
        if stripped.startswith(">"):
            add_callout(doc, stripped[1:].strip(), emphasis=True)
            i += 1
            continue
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            set_para(p, before=0 if not title_done else 18, after=8, line=1.0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, stripped[2:].strip(), size=22 if not title_done else 16, color=INK if not title_done else BLUE)
            for r in p.runs:
                r.bold = True
            title_done = True
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph()
            set_para(p, before=18, after=8, line=1.05)
            add_inline(p, stripped[3:].strip(), size=16, color=BLUE)
            for r in p.runs:
                r.bold = True
            i += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph()
            set_para(p, before=12, after=5, line=1.05)
            add_inline(p, stripped[4:].strip(), size=13, color=DARK_BLUE)
            for r in p.runs:
                r.bold = True
            i += 1
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph()
            set_para(p, before=0, after=4, line=1.12, left=0.30, first=-0.18)
            add_inline(p, "• " + stripped[2:].strip(), size=10.5)
            i += 1
            continue
        if re.match(r"^\d+\. ", stripped):
            p = doc.add_paragraph()
            set_para(p, before=0, after=4, line=1.12, left=0.30, first=-0.18)
            add_inline(p, stripped, size=10.5)
            i += 1
            continue
        if stripped.startswith(">"):
            add_callout(doc, stripped[1:].strip())
            i += 1
            continue
        p = doc.add_paragraph()
        set_para(p, before=0, after=6, line=1.10)
        add_inline(p, stripped, size=10.5)
        i += 1

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: build_tpeng_proposal_doc.py SOURCE.md OUTPUT.docx")
    build(Path(sys.argv[1]), Path(sys.argv[2]))
